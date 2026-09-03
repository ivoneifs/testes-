"""Geração do laudo neuropsicológico integrado em .docx (Word).

Usa apenas os dados já produzidos pelo fluxo (paciente + resultado estruturado da
IA + testes corrigidos). Não chama serviços externos. O texto continua exigindo
revisão e assinatura de profissional habilitado.
"""
from __future__ import annotations

import base64
import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x1B, 0x23, 0x33)
ACCENT = RGBColor(0x2F, 0x57, 0xB8)
MUTED = RGBColor(0x60, 0x6A, 0x80)

SECTION_ORDER = [
    ("identificacao_e_demanda", "Identificação e demanda"),
    ("historia_clinica_e_desenvolvimental", "História clínica e desenvolvimental"),
    ("procedimentos_e_instrumentos", "Procedimentos e instrumentos"),
    ("analise_instrumentos", "Análise de instrumentos e interpretações"),
    ("resultados_por_dominio", "Resultados por domínio"),
    ("integracao_neuropsicologica", "Integração neuropsicológica"),
    ("hipoteses_e_diferenciais", "Hipóteses e diagnósticos diferenciais"),
    ("recomendacoes", "Recomendações"),
    ("limitacoes", "Limitações"),
    ("conclusao", "Conclusão"),
]


def _clean(v: Any) -> str:
    return "" if v is None else str(v).strip()


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Corpo de texto justificado (List Bullet herda de Normal, mas fixa explicitamente).
    for style_name in ("List Bullet",):
        try:
            doc.styles[style_name].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT


def _body(doc: Document, text: str) -> None:
    text = _clean(text)
    if not text:
        return
    for para in text.split("\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)


def _bullets(doc: Document, items: list) -> None:
    for item in items or []:
        item = _clean(item)
        if item:
            doc.add_paragraph(item, style="List Bullet")


def _data_table(doc: Document, columns: list, rows: list) -> None:
    columns = [_clean(c) for c in (columns or []) if _clean(c)]
    rows = [r for r in (rows or []) if isinstance(r, list)]
    if not columns or not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = "Table Grid"
    for j, name in enumerate(columns):
        run = table.rows[0].cells[j].paragraphs[0].add_run(name)
        run.bold = True
        run.font.size = Pt(9)
    for i, row in enumerate(rows, start=1):
        for j in range(len(columns)):
            val = _clean(row[j]) if j < len(row) else ""
            table.rows[i].cells[j].paragraphs[0].add_run(val).font.size = Pt(9)
    doc.add_paragraph()


def _instrument_analysis(doc: Document, entries: list) -> None:
    for entry in entries or []:
        if not isinstance(entry, dict):
            _body(doc, entry)
            continue
        name = _clean(entry.get("instrumento"))
        if name:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(name)
            r.bold = True
            r.font.size = Pt(10.5)
        _body(doc, entry.get("objetivo"))
        for tbl in entry.get("tabelas") or []:
            if not isinstance(tbl, dict):
                continue
            title = _clean(tbl.get("titulo"))
            if title:
                tp = doc.add_paragraph()
                tp.paragraph_format.space_before = Pt(4)
                tr = tp.add_run(title)
                tr.italic = True
                tr.font.size = Pt(9.5)
            _data_table(doc, tbl.get("colunas"), tbl.get("linhas"))
        _body(doc, entry.get("comentario"))


def _charts_section(doc: Document, charts: list) -> None:
    """Embute os gráficos (PNG data URL) enviados pelo frontend, um por teste."""
    ok = False
    for c in charts or []:
        img = _clean(c.get("image"))
        if not img.startswith("data:image"):
            continue
        try:
            data = base64.b64decode(img.split(",", 1)[1])
        except Exception:
            continue
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(6)
        cr = cap.add_run(" — ".join(x for x in (_clean(c.get("test")), _clean(c.get("title"))) if x))
        cr.italic = True
        cr.font.size = Pt(9.5)
        try:
            doc.add_picture(io.BytesIO(data), width=Inches(5.9))
        except Exception:
            continue
        ok = True
    return ok


def _domain_block(doc: Document, entries: list) -> None:
    for entry in entries or []:
        if not isinstance(entry, dict):
            _body(doc, entry)
            continue
        name = _clean(entry.get("dominio"))
        if name:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(name)
            r.bold = True
        _body(doc, entry.get("descricao"))
        ev = [e for e in (entry.get("evidencias") or []) if _clean(e)]
        if ev:
            p = doc.add_paragraph()
            p.add_run("Evidências: ").italic = True
            p.add_run("; ".join(_clean(e) for e in ev))


def build_integrated_docx(patient: dict, report: dict, tests: list[str] | None = None,
                          charts: list | None = None) -> bytes:
    doc = Document()
    _style(doc)

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Pt(64)
    sec.top_margin = sec.bottom_margin = Pt(56)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("AVALIAÇÃO NEUROPSICOLÓGICA COMPLETA")
    tr.bold = True
    tr.font.size = Pt(17)
    tr.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(14)

    # Cabeçalho de identificação
    meta_rows = [
        ("Nome", _clean(patient.get("name")) or "—"),
        ("Data de nascimento", _clean(patient.get("birth_date")) or "—"),
        ("Data de aplicação", _clean(patient.get("application_date")) or "—"),
        ("Sexo", _clean(patient.get("sex")) or "—"),
        ("Escolaridade", _clean(patient.get("education")) or "—"),
        ("Instrumentos aplicados", ", ".join(t for t in (tests or []) if t) or "—"),
        ("Data de emissão", date.today().strftime("%d/%m/%Y")),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for (label, value), row in zip(meta_rows, table.rows):
        c0, c1 = row.cells
        c0.text = ""
        run = c0.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        c1.text = ""
        c1.paragraphs[0].add_run(value).font.size = Pt(10)

    doc.add_paragraph()

    for key, label in SECTION_ORDER:
        if key not in (report or {}):
            continue
        value = report[key]
        _heading(doc, label)
        if key == "analise_instrumentos":
            _instrument_analysis(doc, value)
            if charts:
                _heading(doc, "Gráficos por teste")
                _charts_section(doc, charts)
        elif key == "resultados_por_dominio":
            _domain_block(doc, value)
        elif isinstance(value, list):
            _bullets(doc, value)
        else:
            _body(doc, value)

    if charts and "analise_instrumentos" not in (report or {}):
        _heading(doc, "Gráficos por teste")
        _charts_section(doc, charts)

    # Assinatura + aviso
    doc.add_paragraph()
    doc.add_paragraph()
    sign = doc.add_paragraph("__________________________________________")
    sign.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph()
    cr = cap.add_run("Profissional responsável — assinatura e registro")
    cr.font.size = Pt(9)
    cr.font.color.rgb = MUTED

    disc = doc.add_paragraph()
    dr = disc.add_run(
        "Documento gerado pelo NeuroScore com apoio de inteligência artificial. "
        "O conteúdo exige revisão, validação clínica e assinatura de profissional "
        "habilitado antes de qualquer uso."
    )
    dr.italic = True
    dr.font.size = Pt(8)
    dr.font.color.rgb = MUTED

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
